"""P2.9C2 — Deterministic draft document export (DOCX / PDF).

Pure rendering layer: same canonical input always produces the same content
order, headings and citation strings. Citations come exclusively from the
deterministic server-side renderer (never from model output). No export
artifact is persisted; no paragraph/revision/source text is ever logged.
"""
from __future__ import annotations

import os
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


class DraftExportError(RuntimeError):
    """Fail-closed export error carrying only a sanitized code."""

    def __init__(self, code: str):
        self.code = code
        super().__init__(code)


# Deterministic Turkish headings for the canonical paragraph types.
TURKISH_SECTION_HEADINGS: dict[str, str] = {
    "merci": "MAKAM",
    "taraflar": "TARAFLAR",
    "konu": "KONU",
    "kisa_ozet": "KISA ÖZET",
    "olaylar": "OLAYLAR",
    "hukuki_degerlendirme": "HUKUKİ DEĞERLENDİRME",
    "deliller": "DELİLLER",
    "hukuki_nedenler": "HUKUKİ NEDENLER",
    "sonuc_ve_talep": "SONUÇ VE TALEP",
    "ekler": "EKLER",
    "body": "",
}

DRAFT_TYPE_LABELS: dict[str, str] = {
    "dava_dilekcesi": "Dava Dilekçesi",
    "cevap_dilekcesi": "Cevap Dilekçesi",
    "cevaba_cevap": "Cevaba Cevap Dilekçesi",
    "ikinci_cevap": "İkinci Cevap Dilekçesi",
    "istinaf": "İstinaf Dilekçesi",
    "temyiz": "Temyiz Dilekçesi",
    "ihtiyati_tedbir": "İhtiyati Tedbir Talebi",
    "itiraz": "İtiraz Dilekçesi",
    "beyan": "Beyan Dilekçesi",
    "delil_listesi": "Delil Listesi",
    "ihtarname": "İhtarname",
    "arabuluculuk_basvurusu": "Arabuluculuk Başvurusu",
}


@dataclass(frozen=True)
class ExportParagraph:
    order: int
    heading: str
    text: str
    citations: tuple[str, ...]


@dataclass(frozen=True)
class ExportDocument:
    title: str
    draft_type: str
    draft_type_label: str
    draft_id_short: str
    version: int
    paragraphs: tuple[ExportParagraph, ...]


def export_filename(draft_type: str, draft_id: str, extension: str) -> str:
    """Safe deterministic filename; never user/party/case-fact text."""
    return f"emsalist-{draft_type}-{draft_id[:8]}.{extension}"


# ── DOCX ─────────────────────────────────────────────────────────────────────
def render_docx(document: ExportDocument) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    core = doc.core_properties
    # Safe, deterministic core properties only (no run-varying metadata).
    core.title = document.title
    core.author = "Emsalist"
    core.last_modified_by = "Emsalist"
    core.comments = f"surum {document.version}"

    doc.add_heading(document.draft_type_label, level=1)
    for paragraph in document.paragraphs:
        if paragraph.heading:
            doc.add_heading(paragraph.heading, level=2)
        doc.add_paragraph(paragraph.text)
        for citation in paragraph.citations:
            citation_paragraph = doc.add_paragraph()
            run = citation_paragraph.add_run(f"Kaynak: {citation}")
            run.italic = True
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── PDF (pymupdf; Unicode-safe system font, selectable text) ────────────────
_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "C:\\Windows\\Fonts\\arial.ttf",
    "C:\\Windows\\Fonts\\calibri.ttf",
)

_PAGE_WIDTH = 595.0  # A4 portrait (points)
_PAGE_HEIGHT = 842.0
_MARGIN = 56.7  # 2 cm fixed margins
_BODY_SIZE = 11.0
_HEADING_SIZE = 13.0
_CITATION_SIZE = 9.5
_LINE_FACTOR = 1.45
_FOOTER_SIZE = 9.0


def resolve_pdf_font_path() -> str:
    """Explicitly configured system/Docker font; never a committed font file."""
    override = os.environ.get("EMSALIST_PDF_FONT_PATH", "").strip()
    candidates = ([override] if override else []) + list(_FONT_CANDIDATES)
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            return candidate
    raise DraftExportError("draft_export_font_unavailable")


def _wrap_line(font, text: str, size: float, max_width: float) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if not current or font.text_length(candidate, fontsize=size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def render_pdf(document: ExportDocument) -> bytes:
    import fitz

    font_path = resolve_pdf_font_path()
    font = fitz.Font(fontfile=font_path)
    usable_width = _PAGE_WIDTH - 2 * _MARGIN
    bottom_limit = _PAGE_HEIGHT - _MARGIN - _FOOTER_SIZE * 2

    # Deterministic layout plan: (size, text, extra_gap_before) line entries.
    entries: list[tuple[float, str, float]] = [(_HEADING_SIZE + 2, document.title, 0.0)]
    for paragraph in document.paragraphs:
        if paragraph.heading:
            for line in _wrap_line(font, paragraph.heading, _HEADING_SIZE, usable_width):
                entries.append((_HEADING_SIZE, line, _HEADING_SIZE))
        first = True
        for line in _wrap_line(font, paragraph.text, _BODY_SIZE, usable_width):
            entries.append((_BODY_SIZE, line, _BODY_SIZE * 0.6 if first else 0.0))
            first = False
        for citation in paragraph.citations:
            first = True
            for line in _wrap_line(font, f"Kaynak: {citation}", _CITATION_SIZE,
                                   usable_width):
                entries.append((_CITATION_SIZE, line,
                                _CITATION_SIZE * 0.5 if first else 0.0))
                first = False

    pdf = fitz.open()
    pages: list[list[tuple[float, float, str]]] = [[]]
    cursor = _MARGIN
    for size, line, gap_before in entries:
        line_height = size * _LINE_FACTOR
        if cursor + gap_before + line_height > bottom_limit:
            pages.append([])
            cursor = _MARGIN
            gap_before = 0.0
        cursor += gap_before + line_height
        pages[-1].append((size, cursor, line))

    total_pages = len(pages)
    for page_number, lines in enumerate(pages, start=1):
        page = pdf.new_page(width=_PAGE_WIDTH, height=_PAGE_HEIGHT)
        writer = fitz.TextWriter(page.rect)
        for size, y_position, line in lines:
            if line:
                writer.append((_MARGIN, y_position), line, font=font, fontsize=size)
        footer = f"Sayfa {page_number} / {total_pages}"
        footer_width = font.text_length(footer, fontsize=_FOOTER_SIZE)
        writer.append(((_PAGE_WIDTH - footer_width) / 2, _PAGE_HEIGHT - _MARGIN / 2),
                      footer, font=font, fontsize=_FOOTER_SIZE)
        writer.write_text(page)

    # Deterministic, safe metadata only (no creation/mod dates, no producer).
    pdf.set_metadata({"title": document.title, "author": "Emsalist"})
    pdf.del_xml_metadata()
    content = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return content
